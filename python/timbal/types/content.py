import base64
import io
import json
from ast import literal_eval
from typing import Any, Literal

import pandas as pd
import structlog
from docx import Document
from pydantic import BaseModel

from ..handlers.pdfs import convert_pdf_to_images
from .file import File

logger = structlog.get_logger("timbal.types.content")


AVAILABLE_ENCODINGS = [
    "utf-8",
    "cp1252",
    "iso-8859-1",
    "utf-16",
]


class Content(BaseModel):
    """
    A class representing the content of a chat message.
    """

    type: Literal["text", "file", "tool_use", "tool_result"]


    @staticmethod
    def _parse_tool_use_input(input: Any) -> dict[str, Any]:
        """Aux function to parse tool use input into python objects."""
        input_arguments = {}
        if isinstance(input, dict):
            input_arguments = input
        else:
            try:
                input_arguments = json.loads(input)
            except Exception:
                try:
                    input_arguments = literal_eval(input)
                except Exception:
                    logger.error(
                        f"Both json.loads and literal_eval failed when parsing tool_use input: {input}", 
                        exc_info=True
                    )
        return input_arguments
        

    @classmethod 
    def model_validate(cls, value: Any, *args: Any, **kwargs: Any) -> "Content":
        """Validate and convert input formats into a Content instance."""
        if isinstance(value, Content):
            return value
        elif cls is not Content:
            # cls will be diferent from Content when we call model_validate on one of the subclasses
            return super().model_validate(value, *args, **kwargs)
        elif isinstance(value, str):
            return TextContent(text=value)
        elif isinstance(value, File):
            return FileContent(file=value)
        elif isinstance(value, dict):
            content_type = value.get("type", None)
            if content_type == "text":
                return TextContent(text=value.get("text"))
            elif content_type == "file":
                return FileContent(file=File.validate(value.get("file")))
            elif content_type == "tool_use":
                return ToolUseContent(
                    id=value.get("id"), 
                    name=value.get("name"), 
                    input=cls._parse_tool_use_input(value.get("input")),
                )
            elif content_type == "tool_result":
                tool_result_content = value.get("content")
                if not isinstance(tool_result_content, list):
                    tool_result_content = [tool_result_content]
                return ToolResultContent(
                    id=value.get("id"), 
                    content=[cls.model_validate(item) for item in tool_result_content],
                )
        # By default try to convert whatever python object we have into a string. 
        return TextContent(text=str(value))


class TextContent(Content):
    """
    This class represents a text content in a chat message.
    It also provides methods to convert the text content to the input format required by OpenAI and Anthropic.
    """

    type: Literal["text"] = "text"
    text: str 


    def to_openai_input(self) -> dict[str, Any]:
        """Convert the text content to the input format required by OpenAI."""
        return {
            "type": "text", 
            "text": self.text
        }


    def to_anthropic_input(self) -> dict[str, Any]:
        """Convert the text content to the input format required by Anthropic."""
        return {
            "type": "text", 
            "text": self.text
        }


class FileContent(Content):
    """
    This class represents a file content in a chat message. 
    It also provides methods to convert the file content to the input format required by OpenAI and Anthropic.
    """
    type: Literal["file"] = "file"
    file: File
    # Cached openai and anthropic inputs (some conversions are costly, e.g. audio transcriptions).
    _cached_openai_input: Any | None = None
    _cached_anthropic_input: Any | None = None


    def to_openai_input(self) -> dict[str, Any] | list[dict[str, Any]]:
        """Convert the file content to the input format required by OpenAI."""
        if self._cached_openai_input is not None:
            return self._cached_openai_input

        mime = self.file.__content_type__
        
        # Ensure the file pointer is at the start of the file if we need to read it.
        current_position = self.file.tell()
        if current_position != 0:
            self.file.seek(0)

        if (
            (mime and mime.startswith("text/")) or 
            # Files like .jsonl don't have a mime type.
            (self.file.__source_extension__ in [".json", ".jsonl"])
        ):
            # Attempt to decode the binary content into a string guessing the encoding.
            raw_bytes = self.file.read()
            content = None
            for encoding in AVAILABLE_ENCODINGS:
                try:
                    content = raw_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            if content is None:
                raise ValueError(f"Could not decode file {self.file} with any of: {AVAILABLE_ENCODINGS}")
            openai_input = {
                "type": "text",
                "text": content,
            }
            self._cached_openai_input = openai_input
            return openai_input
        elif self.file.__source_extension__ == ".xlsx":
            df = pd.read_excel(io.BytesIO(self.file.read()))
            openai_input = {
                "type": "text",
                "text": df.to_csv(index=False, header=True, sep=","),
            }
            self._cached_openai_input = openai_input
            return openai_input
        elif self.file.__source_extension__ == ".docx":
            doc = Document(io.BytesIO(self.file.read()))
            text_content = []
            # Process document elements in order to maintain structure
            for element in doc.element.body:
                if element.tag.endswith("p"):
                    paragraph = doc.paragraphs[len([e for e in doc.element.body[:doc.element.body.index(element)] if e.tag.endswith("p")])]
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())
                elif element.tag.endswith("tbl"):
                    table = doc.tables[len([e for e in doc.element.body[:doc.element.body.index(element)] if e.tag.endswith('tbl')])]
                    table_content = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            table_content.append(",".join(row_text))
                    text_content.append("\n".join(table_content))
            text_content = "\n".join(text_content)
            openai_input = {
                "type": "text",
                "text": text_content,
            }
            self._cached_openai_input = openai_input
            return openai_input
        elif mime and mime.startswith("image/"):
            url = self.file.to_data_url()
            return {
                "type": "image_url", 
                "image_url": {"url": url},
            }
        elif mime == "application/pdf":
            # TODO Review openai sdk "file" type
            # OpenAI internally gets an image of each page and complements it with the extracted text.
            # ? For all the use cases we've used, extracting just the images has worked well.
            pages = convert_pdf_to_images(self.file)
            logger.info(
                "convert_pdf_to_images", 
                n_pages=len(pages), 
                description=".to_openai_input() implicitly converting input pdf to images...",
            )
            pages_input = []
            for page in pages:
                url = page.to_data_url()
                pages_input.append({
                    "type": "image_url", 
                    "image_url": {"url": url},
                })
            self._cached_openai_input = pages_input
            return pages_input
        elif mime and mime.startswith("audio/"):
            # Some gpt models accept audio files as input. If the user is using a model that doesn't have this capability, the sdk will raise an error
            if self.file.__source_scheme__ == "data":
                base64_data = self.file.__source__.split(",", 1)[1]
            else:
                base64_data = base64.b64encode(self.file.read()).decode("utf-8")
            if "mp3" not in mime and "wav" not in mime:
                raise ValueError(f"Unsupported audio format: {mime}. Must be one of: mp3, wav")
            return {
                "type": "input_audio",
                "input_audio": {
                    "data": base64_data, 
                    "format": "wav" if "wav" in mime else "mp3",
                },
            }
        raise ValueError(f"Unsupported file {self.file}.")


    def to_anthropic_input(self) -> dict[str, Any] | list[dict[str, Any]]:
        """Convert the file content to the input format required by Anthropic."""
        if self._cached_anthropic_input is not None:
            return self._cached_anthropic_input

        mime = self.file.__content_type__

        # Ensure the file pointer is at the start of the file if we need to read it.
        current_position = self.file.tell()
        if current_position != 0:
            self.file.seek(0)

        if (
            (mime and mime.startswith("text/")) or 
            # Files like .jsonl don't have a mime type.
            (self.file.__source_extension__ in [".json", ".jsonl"])
        ):
            # Attempt to decode the binary content into a string guessing the encoding.
            raw_bytes = self.file.read()
            content = None
            for encoding in AVAILABLE_ENCODINGS:
                try:
                    content = raw_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            if content is None:
                raise ValueError(f"Could not decode file {self.file} with any of: {AVAILABLE_ENCODINGS}")
            anthropic_input = {
                "type": "text",
                "text": content,
            }
            self._cached_anthropic_input = anthropic_input
            return anthropic_input
        elif self.file.__source_extension__ == ".xlsx":
            df = pd.read_excel(io.BytesIO(self.file.read()))
            anthropic_input = {
                "type": "text",
                "text": df.to_csv(index=False, header=True, sep=","),
            }
            self._cached_anthropic_input = anthropic_input
            return anthropic_input
        elif self.file.__source_extension__ == ".docx":
            doc = Document(io.BytesIO(self.file.read()))
            text_content = []
            # Process document elements in order to maintain structure
            for element in doc.element.body:
                if element.tag.endswith("p"):
                    paragraph = doc.paragraphs[len([e for e in doc.element.body[:doc.element.body.index(element)] if e.tag.endswith('p')])]
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())
                elif element.tag.endswith("tbl"):
                    table = doc.tables[len([e for e in doc.element.body[:doc.element.body.index(element)] if e.tag.endswith('tbl')])]
                    table_content = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            table_content.append(",".join(row_text))
                    text_content.append("\n".join(table_content))
            text_content = "\n".join(text_content)
            anthropic_input = {
                "type": "text",
                "text": text_content,
            }
            self._cached_anthropic_input = anthropic_input
            return anthropic_input
        elif mime and mime.startswith("image/"):
            url = self.file.to_data_url()
            if url.startswith("data:"):
                base64_data = url.split(",", 1)[1]
                return {
                    "type": "image", 
                    "source": {
                        "type": "base64",
                        "media_type": mime,
                        "data": base64_data,
                    }
                }
            else:
                return {
                    "type": "image", 
                    "source": {
                        "type": "url",
                        "url": url,
                    }
                }
        elif mime == "application/pdf":
            url = self.file.to_data_url()
            if url.startswith("data:"):
                base64_data = url.split(",", 1)[1]
                return {
                    "type": "document", 
                    "source": {
                        "type": "base64",
                        "media_type": mime,
                        "data": base64_data,
                    }
                }
            else:
                return {
                    "type": "document", 
                    "source": {
                        "type": "url",
                        "url": url,
                    }
                }
        raise ValueError(f"Unsupported file {self.file}.")


class ToolUseContent(Content):
    """
    This class represents a tool use content in a chat message.
    It also provides methods to convert the tool use content to the input format required by OpenAI and Anthropic.
    """

    type: Literal["tool_use"] = "tool_use"
    id: str
    name: str
    input: dict[str, Any]


    def to_openai_input(self) -> dict[str, Any]:
        """Convert the tool use content to the input format required by OpenAI."""
        return {
            "id": self.id,
            "type": "function",
            "function": {
                "arguments": json.dumps(self.input),
                "name": self.name
            }
        }


    def to_anthropic_input(self) -> dict[str, Any]:
        """Convert the tool use content to the input format required by Anthropic."""
        return {
            "type": "tool_use",
            "id": self.id,
            "name": self.name,
            "input": self.input,
        }


class ToolResultContent(Content):
    """
    This class represents a tool result content in a chat message.
    It also provides methods to convert the tool result content to the input format required by OpenAI and Anthropic.
    """

    type: Literal["tool_result"] = "tool_result"
    id: str
    content: list[TextContent | FileContent]


    def to_openai_input(self) -> dict[str, Any]:
        """Convert the tool result content to the input format required by OpenAI."""
        return {
            "role": "tool",
            "content": [item.to_openai_input() for item in self.content],
            "tool_call_id": self.id
        }

    def to_anthropic_input(self) -> dict[str, Any]:
        """Convert the tool result content to the input format required by Anthropic."""
        return {
            "type": "tool_result",
            "tool_use_id": self.id,
            "content": [item.to_anthropic_input() for item in self.content],
        }
