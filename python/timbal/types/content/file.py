import base64
import email
import io
from typing import Any, Literal

# `override` was introduced in Python 3.12; use `typing_extensions` for compatibility with older versions
try:
    from typing import override
except ImportError:
    from typing_extensions import override

import pandas as pd
import structlog
from docx import Document

from ..file import File
from .base import BaseContent
from .text import TextContent

logger = structlog.get_logger("timbal.types.content")


AVAILABLE_ENCODINGS = [
    "utf-8",
    "cp1252",
    "iso-8859-1",
    "utf-16",
]


def pdf_to_images(pdf: File, dpi: int = 200) -> list[File]:
    """Convert a PDF file to a list of images files."""
    import tempfile
    from pathlib import Path

    import fitz
    pdf.seek(0) # Ensure the pointer is at the start of the file
    doc = fitz.Document(stream=pdf.read())
    pages = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        pix = page.get_pixmap(matrix=fitz.Matrix(dpi / 72, dpi / 72))
        # TODO Use File.validate(bytes, {"extension": ".png"})
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            tmp_path = Path(f.name)
        pix.save(tmp_path)
        pix_file = File.validate(tmp_path)
        pages.append(pix_file)
    return pages


def _extract_docx_content(docx: File) -> str:
    doc = Document(io.BytesIO(docx.read()))
    text_content = []
    # Process document elements in order to maintain structure
    for element in doc.element.body:
        if element.tag.endswith("p"):
            paragraph = doc.paragraphs[len([e for e in doc.element.body[:doc.element.body.index(element)] if e.tag.endswith("p")])]
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        elif element.tag.endswith("tbl"):
            table = doc.tables[len([e for e in doc.element.body[:doc.element.body.index(element)] if e.tag.endswith('tbl')])]
            table_content = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_content.append(",".join(row_text))
            text_content.append("\n".join(table_content))
    text_content = "\n".join(text_content)
    return text_content


def _extract_email_body(raw_email_content: str) -> str:
    """Extract only the body content from an EML email, excluding headers and metadata."""
    msg = email.message_from_string(raw_email_content)
    
    # Extract body text, preferring plain text over HTML
    body = ""
    
    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            if content_type == "text/plain":
                payload = part.get_payload(decode=True)
                if payload:
                    try:
                        body = payload.decode(part.get_content_charset() or 'utf-8')
                        break  # Prefer plain text
                    except (UnicodeDecodeError, LookupError):
                        for encoding in AVAILABLE_ENCODINGS:
                            try:
                                body = payload.decode(encoding)
                                break
                            except UnicodeDecodeError:
                                continue
            elif content_type == "text/html" and not body:
                # Fallback to HTML if no plain text found
                payload = part.get_payload(decode=True)
                if payload:
                    try:
                        body = payload.decode(part.get_content_charset() or 'utf-8')
                    except (UnicodeDecodeError, LookupError):
                        for encoding in AVAILABLE_ENCODINGS:
                            try:
                                body = payload.decode(encoding)
                                break
                            except UnicodeDecodeError:
                                continue
    else:
        # Single part message
        payload = msg.get_payload(decode=True)
        if payload:
            try:
                body = payload.decode(msg.get_content_charset() or 'utf-8')
            except (UnicodeDecodeError, LookupError):
                for encoding in AVAILABLE_ENCODINGS:
                    try:
                        body = payload.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
    
    return body.strip() if body else ""


def _extract_email_attachments(raw_email_content: str) -> list[dict[str, Any]]:
    """Extract attachments from an EML email."""
    msg = email.message_from_string(raw_email_content)
    attachments = []
    
    if msg.is_multipart():
        for part in msg.walk():
            # Skip the main message part
            if part.get_content_maintype() == 'multipart':
                continue
                
            # Check if this part has a filename (attachment)
            filename = part.get_filename()
            if filename:
                content_type = part.get_content_type()
                payload = part.get_payload(decode=True)
                content_id = part.get('Content-ID')
                
                if payload:
                    # Create a temporary file-like object for the attachment
                    attachment_data = {
                        'filename': filename,
                        'content_type': content_type,
                        'data': payload,
                        'size': len(payload)
                    }
                    if content_id:
                        # Remove angle brackets from Content-ID
                        attachment_data['content_id'] = content_id.strip('<>')
                    attachments.append(attachment_data)
    
    return attachments


class FileContent(BaseContent):
    """File content type for chat messages."""
    type: Literal["file"] = "file"
    file: File
    # TODO Change this to cached properties
    # Cached openai and anthropic inputs (some conversions are costly, e.g. audio transcriptions).
    _cached_openai_responses_input: Any | None = None
    _cached_openai_chat_completions_input: Any | None = None
    _cached_anthropic_input: Any | None = None

    @override
    def to_openai_responses_input(self, **kwargs: Any) -> dict[str, Any] | list[dict[str, Any]]:
        """See base class."""
        if self._cached_openai_responses_input is not None:
            return self._cached_openai_responses_input

        mime = self.file.__content_type__
        
        # Ensure the file pointer is at the start of the file if we need to read it.
        current_position = self.file.tell()
        if current_position != 0:
            self.file.seek(0)

        if self.file.__source_extension__ == ".xlsx":
            df = pd.read_excel(io.BytesIO(self.file.read()))
            openai_responses_input = {
                "type": "input_text",
                "text": df.to_csv(index=False, header=True, sep=","),
            }
            self._cached_openai_responses_input = openai_responses_input
            return openai_responses_input

        elif self.file.__source_extension__ == ".eml":
            raise NotImplementedError("TODO - Parse EML files into openai responses api.")

        elif self.file.__source_extension__ == ".docx":
            content = _extract_docx_content(self.file)
            openai_responses_input = {
                "type": "input_text",
                "text": content,
            }
            self._cached_openai_responses_input = openai_responses_input
            return openai_responses_input

        elif mime and mime.startswith("image/"):
            url = self.file.to_data_url()
            openai_responses_input = {
                "type": "input_image", 
                "image_url": url,
            }
            self._cached_openai_responses_input = openai_responses_input
            return openai_responses_input

        elif mime == "application/pdf":
            # TODO Review this one
            url = self.file.to_data_url()
            openai_responses_input = {
                "type": "input_file", 
                "filename": self.file.name,
                "file_data": url,
            }
            self._cached_openai_responses_input = openai_responses_input
            return openai_responses_input

        elif mime and mime.startswith("audio/"):
            # Some gpt models accept audio files as input. If the user is using a model that doesn't have this capability, the sdk will raise an error
            if self.file.__source_scheme__ == "data":
                base64_data = self.file.__source__.split(",", 1)[1]
            else:
                base64_data = base64.b64encode(self.file.read()).decode("utf-8")
            if "mp3" not in mime and "wav" not in mime:
                raise ValueError(f"Unsupported audio format: {mime}. Must be one of: mp3, wav")
            return {
                "type": "input_audio",
                "input_audio": {
                    "data": base64_data, 
                    "format": "wav" if "wav" in mime else "mp3",
                },
            }
        
        # Attempt to decode the binary content into a string guessing the encoding.
        raw_bytes = self.file.read()
        content = None
        for encoding in AVAILABLE_ENCODINGS:
            try:
                content = raw_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if content is None:
            raise ValueError(f"Could not decode file {self.file} with any of: {AVAILABLE_ENCODINGS}")
        openai_responses_input = {
            "type": "input_text",
            "text": content,
        }
        self._cached_openai_responses_input = openai_responses_input
        return openai_responses_input


    @override
    def to_openai_chat_completions_input(self, **kwargs: Any) -> dict[str, Any] | list[dict[str, Any]]:
        """See base class."""
        if self._cached_openai_chat_completions_input is not None:
            return self._cached_openai_chat_completions_input

        mime = self.file.__content_type__
        
        # Ensure the file pointer is at the start of the file if we need to read it.
        current_position = self.file.tell()
        if current_position != 0:
            self.file.seek(0)

        if self.file.__source_extension__ == ".xlsx":
            df = pd.read_excel(io.BytesIO(self.file.read()))
            openai_input = {
                "type": "text",
                "text": df.to_csv(index=False, header=True, sep=","),
            }
            self._cached_openai_chat_completions_input = openai_input
            return openai_input

        elif self.file.__source_extension__ == ".docx":
            content = _extract_docx_content(self.file)
            openai_input = {
                "type": "text",
                "text": content,
            }
            self._cached_openai_chat_completions_input = openai_input
            return openai_input

        elif mime and mime.startswith("image/"):
            url = self.file.to_data_url()
            return {
                "type": "image_url", 
                "image_url": {"url": url},
            }

        elif mime == "application/pdf":
            # TODO Review openai sdk "file" type
            # OpenAI internally gets an image of each page and complements it with the extracted text.
            # ? For all the use cases we've used, extracting just the images has worked well.
            pages = pdf_to_images(self.file)
            logger.info(
                "pdf_to_images", 
                n_pages=len(pages), 
                description=".to_openai_chat_completions_input() implicitly converting input pdf to images...",
            )
            pages_input = []
            for page in pages:
                url = page.to_data_url()
                pages_input.append({
                    "type": "image_url", 
                    "image_url": {"url": url},
                })
            self._cached_openai_chat_completions_input = pages_input
            return pages_input

        elif self.file.__source_extension__ == ".eml":
            # EML files are text-based email files - extract body content and attachments
            raw_bytes = self.file.read()
            raw_content = None
            for encoding in AVAILABLE_ENCODINGS:
                try:
                    raw_content = raw_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            if raw_content is None:
                raise ValueError(f"Could not decode EML file {self.file} with any of: {AVAILABLE_ENCODINGS}")
            # Extract email body content and attachments
            body_content = _extract_email_body(raw_content)
            attachments = _extract_email_attachments(raw_content)
            msg = email.message_from_string(raw_content)
            processed_attachments = set()
            attachment_counter = 1
            openai_input = []
            if msg.is_multipart():
                for part in msg.walk():
                    content_id = part.get('Content-ID')
                    if content_id:
                        # Remove angle brackets from Content-ID
                        cid = content_id.strip('<>')
                        # Find matching attachment
                        for i, attachment in enumerate(attachments):
                            if attachment.get('content_id') == cid and i not in processed_attachments:
                                # Replace CID reference with descriptive placeholder
                                filename = attachment['filename']
                                placeholder = f"[File {attachment_counter}: {filename}]"
                                # Replace CID reference in body
                                cid_pattern = f"[cid:{cid}]"
                                body_content = body_content.replace(cid_pattern, placeholder)
                                # Add the attachment as FileContent
                                attachment_file = io.BytesIO(attachment['data'])
                                attachment_file.name = attachment['filename']
                                extension = f".{attachment['filename'].split('.')[-1]}" if '.' in attachment['filename'] else None
                                file_obj = File.validate(attachment_file, info={"extension": extension, "content_type": attachment['content_type']})
                                converted = FileContent(file=file_obj).to_openai_chat_completions_input()
                                if isinstance(converted, list):
                                    openai_input.extend(converted)
                                else:
                                    openai_input.append(converted)
                                processed_attachments.add(i)
                                attachment_counter += 1
                                break
            # Process attachments without CID references (like standalone PDFs)
            for i, attachment in enumerate(attachments):
                if not attachment.get('content_id') and i not in processed_attachments:
                    filename = attachment['filename']
                    placeholder = f"[File {attachment_counter}: {filename}]"
                    # Add placeholder to body content
                    body_content += f"\n\n{placeholder}"
                    # Add the attachment as FileContent
                    attachment_file = io.BytesIO(attachment['data'])
                    attachment_file.name = attachment['filename']
                    extension = f".{attachment['filename'].split('.')[-1]}" if '.' in attachment['filename'] else None
                    file_obj = File.validate(attachment_file, info={"extension": extension, "content_type": attachment['content_type']})
                    converted = FileContent(file=file_obj).to_openai_chat_completions_input()
                    if isinstance(converted, list):
                        openai_input.extend(converted)
                    else:
                        openai_input.append(converted)
                    processed_attachments.add(i)
                    attachment_counter += 1
            openai_input.append(TextContent(text=body_content).to_openai_chat_completions_input())
            self._cached_openai_input = openai_input
            return openai_input

        elif mime and mime.startswith("audio/"):
            # Some gpt models accept audio files as input. If the user is using a model that doesn't have this capability, the sdk will raise an error
            if self.file.__source_scheme__ == "data":
                base64_data = self.file.__source__.split(",", 1)[1]
            else:
                base64_data = base64.b64encode(self.file.read()).decode("utf-8")
            if "mp3" not in mime and "wav" not in mime:
                raise ValueError(f"Unsupported audio format: {mime}. Must be one of: mp3, wav")
            return {
                "type": "input_audio",
                "input_audio": {
                    "data": base64_data, 
                    "format": "wav" if "wav" in mime else "mp3",
                },
            }

        # Attempt to decode the binary content into a string guessing the encoding.
        raw_bytes = self.file.read()
        content = None
        for encoding in AVAILABLE_ENCODINGS:
            try:
                content = raw_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if content is None:
            raise ValueError(f"Could not decode file {self.file} with any of: {AVAILABLE_ENCODINGS}")
        openai_input = {
            "type": "text",
            "text": content,
        }
        self._cached_openai_chat_completions_input = openai_input
        return openai_input


    @override
    def to_anthropic_input(self, **kwargs: Any) -> dict[str, Any] | list[dict[str, Any]]:
        """See base class."""
        if self._cached_anthropic_input is not None:
            return self._cached_anthropic_input

        mime = self.file.__content_type__

        # Ensure the file pointer is at the start of the file if we need to read it.
        current_position = self.file.tell()
        if current_position != 0:
            self.file.seek(0)

        if self.file.__source_extension__ == ".xlsx":
            df = pd.read_excel(io.BytesIO(self.file.read()))
            anthropic_input = {
                "type": "text",
                "text": df.to_csv(index=False, header=True, sep=","),
            }
            self._cached_anthropic_input = anthropic_input
            return anthropic_input

        elif self.file.__source_extension__ == ".docx":
            content = _extract_docx_content(self.file)
            anthropic_input = {
                "type": "text",
                "text": content,
            }
            self._cached_anthropic_input = anthropic_input
            return anthropic_input

        elif mime and mime.startswith("image/"):
            # TODO Review this one
            url = self.file.to_data_url()
            if url.startswith("data:"):
                base64_data = url.split(",", 1)[1]
                return {
                    "type": "image", 
                    "source": {
                        "type": "base64",
                        "media_type": mime,
                        "data": base64_data,
                    }
                }
            else:
                return {
                    "type": "image", 
                    "source": {
                        "type": "url",
                        "url": url,
                    }
                }

        elif mime == "application/pdf":
            # TODO Review this one
            url = self.file.to_data_url()
            if url.startswith("data:"):
                base64_data = url.split(",", 1)[1]
                return {
                    "type": "document", 
                    "source": {
                        "type": "base64",
                        "media_type": mime,
                        "data": base64_data,
                    }
                }
            else:
                return {
                    "type": "document", 
                    "source": {
                        "type": "url",
                        "url": url,
                    }
                }

        elif self.file.__source_extension__ == ".eml":
            # EML files are text-based email files - extract body content and attachments
            raw_bytes = self.file.read()
            raw_content = None
            for encoding in AVAILABLE_ENCODINGS:
                try:
                    raw_content = raw_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            if raw_content is None:
                raise ValueError(f"Could not decode EML file {self.file} with any of: {AVAILABLE_ENCODINGS}")
            # Extract email body content and attachments
            body_content = _extract_email_body(raw_content)
            attachments = _extract_email_attachments(raw_content)
            msg = email.message_from_string(raw_content)
            processed_attachments = set()
            attachment_counter = 1
            anthropic_input = []
            if msg.is_multipart():
                for part in msg.walk():
                    content_id = part.get('Content-ID')
                    if content_id:
                        # Remove angle brackets from Content-ID
                        cid = content_id.strip('<>')
                        # Find matching attachment
                        for i, attachment in enumerate(attachments):
                            if attachment.get('content_id') == cid and i not in processed_attachments:
                                # Replace CID reference with descriptive placeholder
                                filename = attachment['filename']
                                placeholder = f"[File {attachment_counter}: {filename}]"
                                # Replace CID reference in body
                                cid_pattern = f"[cid:{cid}]"
                                body_content = body_content.replace(cid_pattern, placeholder)
                                # Add the attachment as FileContent
                                attachment_file = io.BytesIO(attachment['data'])
                                attachment_file.name = attachment['filename']
                                extension = f".{attachment['filename'].split('.')[-1]}" if '.' in attachment['filename'] else None
                                # Ensure the filename has the correct extension for MIME type detection
                                if extension and not attachment['filename'].endswith(extension):
                                    attachment_file.name = attachment['filename'] + extension
                                file_obj = File.validate(attachment_file, info={"extension": extension, "content_type": attachment['content_type']})
                                converted = FileContent(file=file_obj).to_anthropic_input()
                                if isinstance(converted, list):
                                    anthropic_input.extend(converted)
                                else:
                                    anthropic_input.append(converted)
                                
                                processed_attachments.add(i)
                                attachment_counter += 1
                                break
            # Process attachments without CID references (like standalone PDFs)
            for i, attachment in enumerate(attachments):
                if not attachment.get('content_id') and i not in processed_attachments:
                    filename = attachment['filename']
                    placeholder = f"[File {attachment_counter}: {filename}]"
                    # Add placeholder to body content
                    body_content += f"\n\n{placeholder}"
                    # Add the attachment as FileContent
                    attachment_file = io.BytesIO(attachment['data'])
                    attachment_file.name = attachment['filename']
                    extension = f".{attachment['filename'].split('.')[-1]}" if '.' in attachment['filename'] else None
                    # Ensure the filename has the correct extension for MIME type detection
                    if extension and not attachment['filename'].endswith(extension):
                        attachment_file.name = attachment['filename'] + extension
                    file_obj = File.validate(attachment_file, info={"extension": extension, "content_type": attachment['content_type']})
                    converted = FileContent(file=file_obj).to_anthropic_input()
                    if isinstance(converted, list):
                        anthropic_input.extend(converted)
                    else:
                        anthropic_input.append(converted)
                    processed_attachments.add(i)
                    attachment_counter += 1
            # Add text content at the beginning
            anthropic_input.insert(0, TextContent(text=body_content).to_anthropic_input())
            self._cached_anthropic_input = anthropic_input
            return anthropic_input

        # Attempt to decode the binary content into a string guessing the encoding.
        raw_bytes = self.file.read()
        content = None
        for encoding in AVAILABLE_ENCODINGS:
            try:
                content = raw_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if content is None:
            raise ValueError(f"Could not decode file {self.file} with any of: {AVAILABLE_ENCODINGS}")
        anthropic_input = {
            "type": "text",
            "text": content,
        }
        self._cached_anthropic_input = anthropic_input
        return anthropic_input
