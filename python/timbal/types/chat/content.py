"""
Defines the content types for chat messages in Timbal.

Types:
- TextContent: Plain text messages
- FileContent: File attachments
- ToolUseContent: Request to use a specific tool
- ToolResultContent: Result returned by a tool

All message content must be an instance of one of these types.

Usage:

1. Validating a content:
   >>> text_content = Content.model_validate(text_block)
   >>> tool_use_content = Content.model_validate(tool_use_block)

2. Converting a content to the input format required by OpenAI and Anthropic:
   >>> file_content = FileContent(file=File.validate(data_url))
   >>> file_content.to_openai_input()
   >>> file_content.to_anthropic_input()

"""

import base64
import email
import io
import json
from ast import literal_eval
from typing import Any, Literal

import pandas as pd
import structlog
from anthropic.types import TextBlock as AnthropicTextBlock
from anthropic.types import ToolUseBlock as AnthropicToolUseBlock
from docx import Document

try:
    # In newer OpenAI SDK versions, use the concrete function tool call type
    from openai.types.chat.chat_completion_message_function_tool_call import (
        ChatCompletionMessageFunctionToolCall as OpenAIToolCall,
    )
except ImportError:
    # Fallback for older versions
    from openai.types.chat import ChatCompletionMessageToolCall as OpenAIToolCall
from pydantic import BaseModel

# TODO Add a param in the Agent.__init__() where we can customize this.
from ...steps.openai import stt

# from ...steps.elevenlabs import stt
from ...steps.pdfs import convert_pdf_to_images
from ..file import File

logger = structlog.get_logger("timbal.types.chat.content")


AVAILABLE_ENCODINGS = [
    "utf-8",
    "cp1252",
    "iso-8859-1",
    "utf-16",
]


def _extract_email_body(raw_email_content: str) -> str:
    """Extract only the body content from an EML email, excluding headers and metadata."""
    msg = email.message_from_string(raw_email_content)
    
    # Extract body text, preferring plain text over HTML
    body = ""
    
    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            if content_type == "text/plain":
                payload = part.get_payload(decode=True)
                if payload:
                    try:
                        body = payload.decode(part.get_content_charset() or 'utf-8')
                        break  # Prefer plain text
                    except (UnicodeDecodeError, LookupError):
                        for encoding in AVAILABLE_ENCODINGS:
                            try:
                                body = payload.decode(encoding)
                                break
                            except UnicodeDecodeError:
                                continue
            elif content_type == "text/html" and not body:
                # Fallback to HTML if no plain text found
                payload = part.get_payload(decode=True)
                if payload:
                    try:
                        body = payload.decode(part.get_content_charset() or 'utf-8')
                    except (UnicodeDecodeError, LookupError):
                        for encoding in AVAILABLE_ENCODINGS:
                            try:
                                body = payload.decode(encoding)
                                break
                            except UnicodeDecodeError:
                                continue
    else:
        # Single part message
        payload = msg.get_payload(decode=True)
        if payload:
            try:
                body = payload.decode(msg.get_content_charset() or 'utf-8')
            except (UnicodeDecodeError, LookupError):
                for encoding in AVAILABLE_ENCODINGS:
                    try:
                        body = payload.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
    
    return body.strip() if body else ""


def _extract_email_attachments(raw_email_content: str) -> list[dict[str, Any]]:
    """Extract attachments from an EML email."""
    msg = email.message_from_string(raw_email_content)
    attachments = []
    
    if msg.is_multipart():
        for part in msg.walk():
            # Skip the main message part
            if part.get_content_maintype() == 'multipart':
                continue
                
            # Check if this part has a filename (attachment)
            filename = part.get_filename()
            if filename:
                content_type = part.get_content_type()
                payload = part.get_payload(decode=True)
                content_id = part.get('Content-ID')
                
                if payload:
                    # Create a temporary file-like object for the attachment
                    attachment_data = {
                        'filename': filename,
                        'content_type': content_type,
                        'data': payload,
                        'size': len(payload)
                    }
                    if content_id:
                        # Remove angle brackets from Content-ID
                        attachment_data['content_id'] = content_id.strip('<>')
                    attachments.append(attachment_data)
    
    return attachments


class Content(BaseModel):
    """
    A class representing the content of a chat message.
    """

    type: Literal["text", "file", "tool_use", "tool_result"]


    @staticmethod
    def _parse_openai_function_arguments(arguments: Any) -> dict[str, Any]:
        """Aux function to parse openai tool use arguments into python objects."""
        input_arguments = {}

        if isinstance(arguments, dict):
            input_arguments = arguments
        else:
            try:
                input_arguments = json.loads(arguments)
            except Exception:
                try:
                    input_arguments = literal_eval(arguments)
                except Exception:
                    logger.error(
                        "Both json.loads and literal_eval failed on OpenAI function arguments", 
                        exc_info=True
                    )
        
        return input_arguments
        

    @classmethod 
    def model_validate(cls, value: Any, *args: Any, **kwargs: Any) -> "Content":
        """Validate and convert input formats into a Content instance."""
        # Don't recurse if we're already dealing with a Content instance
        if isinstance(value, Content):
            return value
        
        # cls will be diferent from Content when we call model_validate on one of the subclasses
        if cls is not Content:
            return super().model_validate(value, *args, **kwargs)
        
        if isinstance(value, str):
            return TextContent(text=value)

        if isinstance(value, File):
            return FileContent(file=value)
        
        if isinstance(value, AnthropicTextBlock):
            return TextContent(text=value.text)

        if isinstance(value, AnthropicToolUseBlock):
            return ToolUseContent(
                id=value.id,
                name=value.name,
                input=value.input,
            )

        if isinstance(value, OpenAIToolCall):
            arguments = value.function.arguments
            input_arguments = Content._parse_openai_function_arguments(arguments)
            return ToolUseContent(
                id=value.id,
                name=value.function.name,
                input=input_arguments
            )
        
        # TODO Review
        if isinstance(value, dict):
            content_type = value.get("type", None)

            if content_type == "text":
                return TextContent(text=value.get("text"))
            
            # Anthropic's file content type.
            if content_type == "file":
                return FileContent(file=File.validate(value.get("file")))
            
            # OpenAI's file content type.
            if content_type == "image_url":
                return FileContent(file=File.validate(value.get("image_url")['url']))
            
            if content_type == "input_audio":
                return FileContent(file=File.validate(value.get("input_audio")['data']))

            # Anthropic's tool use content.
            if content_type == "tool_use":
                input_value = value.get("input") 
                if isinstance(input_value, str):
                    if input_value != "":
                        input_value = json.loads(input_value)
                    else:
                        input_value = {}
                return ToolUseContent(
                    id=value.get("id"), 
                    name=value.get("name"), 
                    input=input_value,
                )
            
            # OpenAI's tool use content.
            if content_type == "function":
                arguments = value["function"]["arguments"]
                input_arguments = Content._parse_openai_function_arguments(arguments)
                return ToolUseContent(
                    id=value.get("id"),  
                    name=value["function"]["name"],
                    input=input_arguments,
                )
            
            # Anthropic's tool result content.
            if content_type == "tool_result":
                tool_result_content = value.get("content", [])
                if not isinstance(tool_result_content, list):
                    tool_result_content = [tool_result_content]
                return ToolResultContent(
                    id=value.get("tool_use_id") or value.get("id"), 
                    content=[cls.model_validate(item) for item in tool_result_content],
                )
        
        # By default try to convert whatever python object we have into a string. 
        return TextContent(text=str(value))


class FileContent(Content):
    """
    This class represents a file content in a chat message. 
    It also provides methods to convert the file content to the input format required by OpenAI and Anthropic.
    """
    type: Literal["file"] = "file"
    file: File
    # Cached openai and anthropic inputs (some conversions are costly, e.g. audio transcriptions).
    _cached_openai_input: Any | None = None
    _cached_anthropic_input: Any | None = None


    async def to_openai_input(self, model: str | None = None) -> dict[str, Any] | list[dict[str, Any]]:
        """Convert the file content to the input format required by OpenAI."""
        if self._cached_openai_input is not None:
            return self._cached_openai_input

        mime = self.file.__content_type__
        
        # Ensure the file pointer is at the start of the file if we need to read it.
        current_position = self.file.tell()
        if current_position != 0:
            self.file.seek(0)

        if (
            (mime and mime.startswith("text/")) or 
            # Files like .jsonl don't have a mime type.
            (self.file.__source_extension__ in [".json", ".jsonl"])
        ):
            # Attempt to decode the binary content into a string guessing the encoding.
            raw_bytes = self.file.read()

            content = None
            for encoding in AVAILABLE_ENCODINGS:
                try:
                    content = raw_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError(f"Could not decode file {self.file} with any of: {AVAILABLE_ENCODINGS}")

            openai_input = {
                "type": "text",
                "text": content,
            }

            self._cached_openai_input = openai_input
            return openai_input
        
        elif self.file.__source_extension__ == ".xlsx":
            df = pd.read_excel(io.BytesIO(self.file.read()))

            openai_input = {
                "type": "text",
                "text": df.to_csv(index=False, header=True, sep=","),
            }

            self._cached_openai_input = openai_input
            return openai_input

        elif self.file.__source_extension__ == ".docx":
            doc = Document(io.BytesIO(self.file.read()))
            text_content = []
            
            # Process document elements in order to maintain structure
            for element in doc.element.body:

                if element.tag.endswith("p"):
                    paragraph = doc.paragraphs[len([e for e in doc.element.body[:doc.element.body.index(element)] if e.tag.endswith("p")])]
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())

                elif element.tag.endswith("tbl"):
                    table = doc.tables[len([e for e in doc.element.body[:doc.element.body.index(element)] if e.tag.endswith('tbl')])]
                    table_content = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            table_content.append(",".join(row_text))
                    text_content.append("\n".join(table_content))

            text_content = "\n".join(text_content)

            openai_input = {
                "type": "text",
                "text": text_content,
            }
            
            self._cached_openai_input = openai_input
            return openai_input

        elif mime and mime.startswith("image/"):
            url = self.file.to_data_url()
            return {
                "type": "image_url", 
                "image_url": {"url": url},
            }

        elif mime == "application/pdf":
            # ! OpenAI API is broken. They state you can upload pdfs as base64 encoded data, however 
            # it errors with missing_required_parameter 'file_id'. We don't want to upload the files to 
            # the openai platform (as of yet), since we don't want to control org limits and storage.
            # # We need the base64 data of the pdf.
            # data_url = self.file.to_data_url()
            # base64_data = data_url.split(",", 1)[1]
            # return {
            #     "type": "file",
            #     "file": {
            #         "file_name": f"{uuid7()}.pdf",
            #         "file_data": base64_data,
            #     },
            # }
            # OpenAI internally gets an image of each page and complements it with the extracted text.
            # ? For all the use cases we've used, extracting just the images has worked well.
            pages = convert_pdf_to_images(self.file)
            logger.info(
                "convert_pdf_to_images", 
                n_pages=len(pages), 
                description=".to_openai_input() implicitly converting input pdf to images...",
            )
            pages_input = []
            for page in pages:
                url = page.to_data_url()
                pages_input.append({
                    "type": "image_url", 
                    "image_url": {"url": url},
                })

            self._cached_openai_input = pages_input
            return pages_input

        elif self.file.__source_extension__ == ".eml":
            # EML files are text-based email files - extract body content and attachments
            raw_bytes = self.file.read()
            
            raw_content = None
            for encoding in AVAILABLE_ENCODINGS:
                try:
                    raw_content = raw_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if raw_content is None:
                raise ValueError(f"Could not decode EML file {self.file} with any of: {AVAILABLE_ENCODINGS}")

            # Extract email body content and attachments
            body_content = _extract_email_body(raw_content)
            attachments = _extract_email_attachments(raw_content)
            
            msg = email.message_from_string(raw_content)

            processed_attachments = set()
            attachment_counter = 1
            openai_input = []
            
            if msg.is_multipart():
                for part in msg.walk():
                    content_id = part.get('Content-ID')
                    if content_id:
                        # Remove angle brackets from Content-ID
                        cid = content_id.strip('<>')
                        # Find matching attachment
                        for i, attachment in enumerate(attachments):
                            if attachment.get('content_id') == cid and i not in processed_attachments:
                                # Replace CID reference with descriptive placeholder
                                filename = attachment['filename']
                                
                                placeholder = f"[File {attachment_counter}: {filename}]"
                                
                                # Replace CID reference in body
                                cid_pattern = f"[cid:{cid}]"
                                body_content = body_content.replace(cid_pattern, placeholder)
                                
                                # Add the attachment as FileContent
                                attachment_file = io.BytesIO(attachment['data'])
                                attachment_file.name = attachment['filename']
                                extension = f".{attachment['filename'].split('.')[-1]}" if '.' in attachment['filename'] else None
                                file_obj = File.validate(attachment_file, info={"extension": extension, "content_type": attachment['content_type']})
                                converted = await FileContent(file=file_obj).to_openai_input()
                                if isinstance(converted, list):
                                    openai_input.extend(converted)
                                else:
                                    openai_input.append(converted)
                                
                                processed_attachments.add(i)
                                attachment_counter += 1
                                break
            
            # Process attachments without CID references (like standalone PDFs)
            for i, attachment in enumerate(attachments):
                if not attachment.get('content_id') and i not in processed_attachments:
                    filename = attachment['filename']
                    placeholder = f"[File {attachment_counter}: {filename}]"
                    
                    # Add placeholder to body content
                    body_content += f"\n\n{placeholder}"
                    
                    # Add the attachment as FileContent
                    attachment_file = io.BytesIO(attachment['data'])
                    attachment_file.name = attachment['filename']
                    extension = f".{attachment['filename'].split('.')[-1]}" if '.' in attachment['filename'] else None
                    file_obj = File.validate(attachment_file, info={"extension": extension, "content_type": attachment['content_type']})
                    converted = await FileContent(file=file_obj).to_openai_input()
                    if isinstance(converted, list):
                        openai_input.extend(converted)
                    else:
                        openai_input.append(converted)
                    
                    processed_attachments.add(i)
                    attachment_counter += 1

            openai_input.append(await TextContent(text=body_content).to_openai_input())

            self._cached_openai_input = openai_input
            return openai_input

        elif mime and mime.startswith("audio/"):
            gpt_audio_models = [
                "gpt-4o-audio-preview", "gpt-4o-mini-audio-preview", 
                "gpt-4o-realtime-preview", "gpt-4o-mini-realtime-preview",
            ]
            if model in gpt_audio_models or model.startswith("gemini"):
                if self.file.__source_scheme__ == "data":
                    base64_data = self.file.__source__.split(",", 1)[1]
                else:
                    base64_data = base64.b64encode(self.file.read()).decode("utf-8")

                if "mp3" not in mime and "wav" not in mime:
                    raise ValueError(f"Unsupported audio format: {mime}. Must be one of: mp3, wav")

                return {
                    "type": "input_audio",
                    "input_audio": {
                        "data": base64_data, 
                        "format": "wav" if "wav" in mime else "mp3",
                    },
                }
            else:
                # TODO Add the cost of this operation.
                transcription = await stt(audio_file=self.file)
                logger.info(
                    "stt", 
                    transcription=transcription,
                    description=".to_openai_input() implicitly transcribed audio to text..."
                )
                openai_input = {
                    "type": "text",
                    "text": transcription
                }
                self._cached_openai_input = openai_input
                return openai_input

        raise ValueError(f"Unsupported file {self.file}.")


    async def to_anthropic_input(self, model: str | None = None) -> dict[str, Any] | list[dict[str, Any]]:
        """Convert the file content to the input format required by Anthropic."""
        if self._cached_anthropic_input is not None:
            return self._cached_anthropic_input

        mime = self.file.__content_type__

        # Ensure the file pointer is at the start of the file if we need to read it.
        current_position = self.file.tell()
        if current_position != 0:
            self.file.seek(0)

        if (
            (mime and mime.startswith("text/")) or 
            # Files like .jsonl don't have a mime type.
            (self.file.__source_extension__ in [".json", ".jsonl"])
        ):
            # Attempt to decode the binary content into a string guessing the encoding.
            raw_bytes = self.file.read()

            content = None
            for encoding in AVAILABLE_ENCODINGS:
                try:
                    content = raw_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError(f"Could not decode file {self.file} with any of: {AVAILABLE_ENCODINGS}")

            anthropic_input = {
                "type": "text",
                "text": content,
            }

            self._cached_anthropic_input = anthropic_input
            return anthropic_input
        
        elif self.file.__source_extension__ == ".xlsx":
            df = pd.read_excel(io.BytesIO(self.file.read()))

            anthropic_input = {
                "type": "text",
                "text": df.to_csv(index=False, header=True, sep=","),
            }

            self._cached_anthropic_input = anthropic_input
            return anthropic_input

        elif self.file.__source_extension__ == ".docx":
            doc = Document(io.BytesIO(self.file.read()))
            text_content = []
            
            # Process document elements in order to maintain structure
            for element in doc.element.body:

                if element.tag.endswith("p"):
                    paragraph = doc.paragraphs[len([e for e in doc.element.body[:doc.element.body.index(element)] if e.tag.endswith('p')])]
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())

                elif element.tag.endswith("tbl"):
                    table = doc.tables[len([e for e in doc.element.body[:doc.element.body.index(element)] if e.tag.endswith('tbl')])]
                    table_content = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            table_content.append(",".join(row_text))
                    text_content.append("\n".join(table_content))

            text_content = "\n".join(text_content)

            anthropic_input = {
                "type": "text",
                "text": text_content,
            }

            self._cached_anthropic_input = anthropic_input
            return anthropic_input

        elif mime and mime.startswith("image/"):
            url = self.file.to_data_url()
            if url.startswith("data:"):
                base64_data = url.split(",", 1)[1]
                return {
                    "type": "image", 
                    "source": {
                        "type": "base64",
                        "media_type": mime,
                        "data": base64_data,
                    }
                }
            else:
                return {
                    "type": "image", 
                    "source": {
                        "type": "url",
                        "url": url,
                    }
                }

        elif mime == "application/pdf":
            url = self.file.to_data_url()
            if url.startswith("data:"):
                base64_data = url.split(",", 1)[1]
                return {
                    "type": "document", 
                    "source": {
                        "type": "base64",
                        "media_type": mime,
                        "data": base64_data,
                    }
                }
            else:
                return {
                    "type": "document", 
                    "source": {
                        "type": "url",
                        "url": url,
                    }
                }
        
        elif self.file.__source_extension__ == ".eml":
            # EML files are text-based email files - extract body content and attachments
            raw_bytes = self.file.read()
            
            raw_content = None
            for encoding in AVAILABLE_ENCODINGS:
                try:
                    raw_content = raw_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if raw_content is None:
                raise ValueError(f"Could not decode EML file {self.file} with any of: {AVAILABLE_ENCODINGS}")

            # Extract email body content and attachments
            body_content = _extract_email_body(raw_content)
            attachments = _extract_email_attachments(raw_content)
            
            msg = email.message_from_string(raw_content)

            processed_attachments = set()
            attachment_counter = 1
            anthropic_input = []
            
            if msg.is_multipart():
                for part in msg.walk():
                    content_id = part.get('Content-ID')
                    if content_id:
                        # Remove angle brackets from Content-ID
                        cid = content_id.strip('<>')
                        # Find matching attachment
                        for i, attachment in enumerate(attachments):
                            if attachment.get('content_id') == cid and i not in processed_attachments:
                                # Replace CID reference with descriptive placeholder
                                filename = attachment['filename']
                                
                                placeholder = f"[File {attachment_counter}: {filename}]"
                                
                                # Replace CID reference in body
                                cid_pattern = f"[cid:{cid}]"
                                body_content = body_content.replace(cid_pattern, placeholder)
                                
                                # Add the attachment as FileContent
                                attachment_file = io.BytesIO(attachment['data'])
                                attachment_file.name = attachment['filename']
                                extension = f".{attachment['filename'].split('.')[-1]}" if '.' in attachment['filename'] else None
                                # Ensure the filename has the correct extension for MIME type detection
                                if extension and not attachment['filename'].endswith(extension):
                                    attachment_file.name = attachment['filename'] + extension
                                file_obj = File.validate(attachment_file, info={"extension": extension, "content_type": attachment['content_type']})
                                converted = await FileContent(file=file_obj).to_anthropic_input()
                                if isinstance(converted, list):
                                    anthropic_input.extend(converted)
                                else:
                                    anthropic_input.append(converted)
                                
                                processed_attachments.add(i)
                                attachment_counter += 1
                                break
            
            # Process attachments without CID references (like standalone PDFs)
            for i, attachment in enumerate(attachments):
                if not attachment.get('content_id') and i not in processed_attachments:
                    filename = attachment['filename']
                    placeholder = f"[File {attachment_counter}: {filename}]"
                    
                    # Add placeholder to body content
                    body_content += f"\n\n{placeholder}"
                    
                    # Add the attachment as FileContent
                    attachment_file = io.BytesIO(attachment['data'])
                    attachment_file.name = attachment['filename']
                    extension = f".{attachment['filename'].split('.')[-1]}" if '.' in attachment['filename'] else None
                    # Ensure the filename has the correct extension for MIME type detection
                    if extension and not attachment['filename'].endswith(extension):
                        attachment_file.name = attachment['filename'] + extension
                    file_obj = File.validate(attachment_file, info={"extension": extension, "content_type": attachment['content_type']})
                    converted = await FileContent(file=file_obj).to_anthropic_input()
                    if isinstance(converted, list):
                        anthropic_input.extend(converted)
                    else:
                        anthropic_input.append(converted)
                    
                    processed_attachments.add(i)
                    attachment_counter += 1

            # Add text content at the beginning
            anthropic_input.insert(0, await TextContent(text=body_content).to_anthropic_input())

            self._cached_anthropic_input = anthropic_input
            return anthropic_input

        elif mime and mime.startswith("audio/"):
            # TODO Add the cost of this operation.
            transcription = await stt(audio_file=self.file)
            logger.info(
                "stt", 
                transcription=transcription,
                description=".to_anthropic_input() implicitly transcribed audio to text..."
            )

            anthropic_input = {
                "type": "text",
                "text": transcription
            }

            self._cached_anthropic_input = anthropic_input
            return anthropic_input
            
        raise ValueError(f"Unsupported file {self.file}.")


class TextContent(Content):
    """
    This class represents a text content in a chat message.
    It also provides methods to convert the text content to the input format required by OpenAI and Anthropic.
    """

    type: Literal["text"] = "text"
    text: str 


    async def to_openai_input(self, model: str | None = None) -> dict[str, Any]:
        """Convert the text content to the input format required by OpenAI."""
        return {
            "type": "text", 
            "text": self.text
        }


    async def to_anthropic_input(self, model: str | None = None) -> dict[str, Any]:
        """Convert the text content to the input format required by Anthropic."""
        return {
            "type": "text", 
            "text": self.text
        }


class ToolUseContent(Content):
    """
    This class represents a tool use content in a chat message.
    It also provides methods to convert the tool use content to the input format required by OpenAI and Anthropic.
    """

    type: Literal["tool_use"] = "tool_use"
    id: str
    name: str
    input: dict[str, Any]


    async def to_openai_input(self, model: str | None = None) -> dict[str, Any]:
        """Convert the tool use content to the input format required by OpenAI."""
        return {
            "id": self.id,
            "type": "function",
            "function": {
                "arguments": json.dumps(self.input),
                "name": self.name
            }
        }


    async def to_anthropic_input(self, model: str | None = None) -> dict[str, Any]:
        """Convert the tool use content to the input format required by Anthropic."""
        return {
            "type": "tool_use",
            "id": self.id,
            "name": self.name,
            "input": self.input,
        }


class ToolResultContent(Content):
    """
    This class represents a tool result content in a chat message.
    It also provides methods to convert the tool result content to the input format required by OpenAI and Anthropic.
    """

    type: Literal["tool_result"] = "tool_result"
    id: str
    content: list[TextContent | FileContent]


    async def to_openai_input(self, model: str | None = None) -> dict[str, Any]:
        """Convert the tool result content to the input format required by OpenAI."""
        return {
            "role": "tool",
            "content": [await item.to_openai_input(model=model) for item in self.content],
            "tool_call_id": self.id
        }

    async def to_anthropic_input(self, model: str | None = None) -> dict[str, Any]:
        """Convert the tool result content to the input format required by Anthropic."""
        return {
            "type": "tool_result",
            "tool_use_id": self.id,
            "content": [await item.to_anthropic_input(model=model) for item in self.content],
        }
